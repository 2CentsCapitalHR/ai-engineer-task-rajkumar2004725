from docx import Document as DocxDocument
import docx2txt
import re
import os
from rag_pipeline import rag_chain

def parse_docx(file_path):
    """Extract text and metadata from a .docx file."""
    try:
        text = docx2txt.process(file_path)
        doc = DocxDocument(file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        return {
            "text": text,
            "filename": file_path.split('\\')[-1],
            "headings": headings
        }
    except Exception as e:
        print(f"Error parsing {file_path}: {e}")
        return None

def identify_document_type(parsed_doc, rag_chain):
    """Identify document type using filename, headings, and RAG."""
    if not parsed_doc:
        return "Unknown"
    
    filename = parsed_doc["filename"].lower()
    text = parsed_doc["text"]
    if "articles of association" in filename or "aoa" in filename:
        return "Articles of Association (AoA)"
    if "memorandum of association" in filename or "moa" in filename:
        return "Memorandum of Association (MoA/MoU)"
    if "ubo declaration" in filename:
        return "UBO Declaration Form"
    if "employment contract" in filename:
        return "Standard Employment Contract"
    
    query = f"Based on this content, what type of ADGM document is this? Content: {text[:500]}"
    result = rag_chain.invoke({"query": query})
    return result["result"].strip() or "Unknown"

def detect_red_flags(parsed_doc, rag_chain):
    """Detect legal red flags in the document text."""
    if not parsed_doc:
        return []
    
    text = parsed_doc["text"].lower()
    red_flags = []
    
    if "uae federal courts" in text or "federal courts" in text:
        red_flags.append({
            "issue": "Incorrect jurisdiction referenced (e.g., UAE Federal Courts). ADGM requires ADGM Courts.",
            "section": "Jurisdiction Clause",
            "severity": "High",
            "suggestion": "Replace with 'ADGM Courts' per ADGM Companies Regulations 2020, Art. 6."
        })
    
    if "signature" not in text and "signed" not in text:
        red_flags.append({
            "issue": "Missing signatory section or improper formatting.",
            "section": "Signatory Section",
            "severity": "High",
            "suggestion": "Add a signed section with all required parties per ADGM template guidelines."
        })


    print("testing rag pipeline for red flags")
    query = f"Based on ADGM regulations, are there any invalid clauses in this text? Text: {text[:1000]}"
    result = rag_chain.invoke({"query": query})
    if "invalid" in result["result"].lower() or "non-compliant" in result["result"].lower():
        red_flags.append({
            "issue": result["result"].strip(),
            "section": "General Clauses",
            "severity": "Medium",
            "suggestion": "Review and align with ADGM templates."
        })
    
    return red_flags

def add_inline_comments(file_path, red_flags):
    """Add text annotations for red flags at the end of the .docx file."""
    doc = DocxDocument(file_path)
    if not red_flags:
        print("No red flags to comment on.")
        return file_path  # Return original if no changes
    
    # Add a "Comments" section at the end
    comments_section = doc.add_paragraph()
    comments_section.add_run("=== Comments Section ===").bold = True
    comments_section.add_run("\nThe following issues were detected:\n")

    for i, flag in enumerate(red_flags, 1):
        comment_text = f"\n{i}. **Section:** {flag['section']}\n   - **Issue:** {flag['issue']}\n   - **Severity:** {flag['severity']}\n   - **Suggestion:** {flag['suggestion']}\n"
        comments_section.add_run(comment_text)

    output_dir = "reviewed_samples/"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"reviewed_{os.path.basename(file_path)}")
    doc.save(output_path)
    print(f"Saved reviewed file with annotations to: {output_path}")
    return output_path

# Test block
if __name__ == "__main__":
    sample_file = r"samples\sample_aoa.docx"
    parsed = parse_docx(sample_file)
    if parsed:
        doc_type = identify_document_type(parsed, rag_chain)
        red_flags = detect_red_flags(parsed, rag_chain)
        print(f"Parsed document: {parsed}")
        print(f"Identified document type: {doc_type}")
        print(f"Red flags detected: {red_flags}")
        if red_flags:
            reviewed_file = add_inline_comments(sample_file, red_flags)
            print(f"Reviewed file saved as: {reviewed_file}")
    else:
        print("Failed to parse the sample document.")